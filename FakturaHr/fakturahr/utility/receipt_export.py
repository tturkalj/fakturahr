# -*- coding: utf-8 -*-
import os
from docx import Document
from fakturahr import get_root_path
from fakturahr.models.models import ReceiptItem


def get_receipt_document(receipt):
    template_path = os.path.join(get_root_path(), 'utility', 'receipt_template.docx')
    receipt_document = Document(template_path)
    styles = receipt_document.styles
    table_contents_style = styles['Table Contents']
    table_contents_small_style = styles['Table Contents Small']

    client_table = receipt_document.tables[1]
    client_data_cell = client_table.rows[0].cells[0]
    client_data_cell_p = client_data_cell.paragraphs[0]
    client_data_cell_p.add_run(receipt.client.get_name())
    client_data_cell.add_paragraph(receipt.client.get_address()).style = table_contents_style
    client_address = u'{0} {1}'.format(receipt.client.get_postal_code(), receipt.client.get_city())
    client_data_cell.add_paragraph(client_address).style = table_contents_style
    client_data_cell.add_paragraph(u'OIB: {0}'.format(receipt.client.get_oib())).style=table_contents_style

    receipt_data_cell = client_table.rows[0].cells[1]
    receipt_data_cell_p = receipt_data_cell.paragraphs[0]
    receipt_data_cell_p.add_run(u'Broj računa: {0}'.format(receipt.get_number()))
    receipt_data_cell.add_paragraph(u'Datum izdavanja: {0}'.format(receipt.get_issued_date())).style = table_contents_style
    receipt_data_cell.add_paragraph(u'Datum valute: {0}'.format(receipt.get_currency_date())).style = table_contents_style
    receipt_data_cell.add_paragraph(u'Vrijeme izdavanja: {0}'.format(receipt.get_issued_time())).style = table_contents_style
    receipt_data_cell.add_paragraph(u'Mjesto izdavanja: {0}'.format(receipt.get_issued_location())).style = table_contents_style
    receipt_data_cell.add_paragraph(u'Način plaćanja: {0}'.format(receipt.get_payment_type())).style = table_contents_style
    receipt_data_cell.add_paragraph(u'Operater: {0}'.format(receipt.get_operator())).style = table_contents_style

    receipt_items_table = receipt_document.tables[2]
    receipt_items = receipt.items.filter(ReceiptItem.deleted == False).order_by(ReceiptItem.id.asc())
    for index, item in enumerate(receipt_items, 1):
        new_row = receipt_items_table.add_row()
        new_row.cells[0].paragraphs[0].text = u'{0}'.format(index)
        new_row.cells[0].paragraphs[0].style = table_contents_small_style

        new_row.cells[1].paragraphs[0].text= u'{0}'.format(item.get_ean())
        new_row.cells[1].paragraphs[0].style = table_contents_small_style

        new_row.cells[2].paragraphs[0].text = u'{0}'.format(item.get_name())
        new_row.cells[2].paragraphs[0].style = table_contents_small_style

        new_row.cells[3].paragraphs[0].text = u'{0}'.format(item.get_measurement_unit())
        new_row.cells[3].paragraphs[0].style = table_contents_small_style

        new_row.cells[4].paragraphs[0].text = u'{0}'.format(item.get_quantity())
        new_row.cells[4].paragraphs[0].style = table_contents_small_style

        new_row.cells[5].paragraphs[0].text = u'{0}'.format(item.get_item_price())
        new_row.cells[5].paragraphs[0].style = table_contents_small_style

        new_row.cells[6].paragraphs[0].text = u'{0}'.format(item.get_rebate_percent())
        new_row.cells[6].paragraphs[0].style = table_contents_small_style

        new_row.cells[7].paragraphs[0].text = u'{0}'.format(item.get_item_price_sum())
        new_row.cells[7].paragraphs[0].style = table_contents_small_style

    totals_table = receipt_document.tables[3]
    totals_table_rows = totals_table.rows
    for row_index, row in enumerate(totals_table_rows):
        if row_index == 0:
            row.cells[2].paragraphs[0].text = receipt.get_base_amount_formatted()
        elif row_index == 1:
            row.cells[2].paragraphs[0].text = receipt.get_tax_amount_formatted()
        elif row_index == 2:
            row.cells[2].paragraphs[0].text = receipt.get_return_amount_formatted()
        else:
            row.cells[2].paragraphs[0].text = receipt.get_total_amount_formatted()
        row.cells[2].paragraphs[0].style = table_contents_style

    return receipt_document
